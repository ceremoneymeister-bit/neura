#!/usr/bin/env python3
"""DOCX generator for Neura v2 capsules.

Generates professional A4 Word documents from Markdown content
following Russian business document standards (ГОСТ Р 7.0.97-2016).

Usage:
    python3 tools/docx-generator.py --input /tmp/content.md --output /tmp/result.docx
    python3 tools/docx-generator.py --input /tmp/kp.md --output /tmp/kp.docx --template kp
    python3 tools/docx-generator.py --input /tmp/letter.md --output /tmp/letter.docx --template letter

Templates:
    plain   — стандартный деловой документ (default)
    kp      — коммерческое предложение
    letter  — деловое письмо
"""

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.section import WD_ORIENT
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


# ── ГОСТ Р 7.0.97-2016 defaults ────────────────────────────────
MARGIN_TOP = Cm(2.0)
MARGIN_BOTTOM = Cm(2.0)
MARGIN_LEFT = Cm(3.0)
MARGIN_RIGHT = Cm(1.5)
FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)
LINE_SPACING = 1.5
FIRST_LINE_INDENT = Cm(1.25)
HEADING1_SIZE = Pt(16)
HEADING2_SIZE = Pt(14)
HEADING3_SIZE = Pt(13)


def setup_document(template: str = "plain") -> Document:
    """Create a new document with proper page setup."""
    doc = Document()

    # Page setup — A4
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT

    # Default paragraph style
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.first_line_indent = FIRST_LINE_INDENT

    # Heading styles
    for level, size in [(1, HEADING1_SIZE), (2, HEADING2_SIZE), (3, HEADING3_SIZE)]:
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = FONT_NAME
        hs.font.size = size
        hs.font.bold = True
        hs.font.color.rgb = RGBColor(0, 0, 0)
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)
        hs.paragraph_format.first_line_indent = Cm(0)

    return doc


def parse_markdown(text: str) -> list[dict]:
    """Parse simple markdown into structured elements."""
    elements = []
    lines = text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # Headings
        if line.startswith("### "):
            elements.append({"type": "heading", "level": 3, "text": line[4:].strip()})
        elif line.startswith("## "):
            elements.append({"type": "heading", "level": 2, "text": line[3:].strip()})
        elif line.startswith("# "):
            elements.append({"type": "heading", "level": 1, "text": line[2:].strip()})
        # Bullet lists
        elif line.strip().startswith("- ") or line.strip().startswith("• "):
            text_content = re.sub(r"^[\s]*[-•]\s*", "", line)
            elements.append({"type": "bullet", "text": text_content})
        # Numbered lists
        elif re.match(r"^\s*\d+[.)]\s", line):
            text_content = re.sub(r"^\s*\d+[.)]\s*", "", line)
            elements.append({"type": "numbered", "text": text_content})
        # Horizontal rule
        elif line.strip() in ("---", "***", "___"):
            elements.append({"type": "hr"})
        # Table (pipe-delimited)
        elif "|" in line and i + 1 < len(lines) and "---" in lines[i + 1]:
            headers = [c.strip() for c in line.split("|") if c.strip()]
            i += 1  # skip separator
            rows = []
            while i + 1 < len(lines) and "|" in lines[i + 1]:
                i += 1
                row = [c.strip() for c in lines[i].split("|") if c.strip()]
                rows.append(row)
            elements.append({"type": "table", "headers": headers, "rows": rows})
        # Empty line
        elif not line.strip():
            elements.append({"type": "empty"})
        # Regular paragraph
        else:
            elements.append({"type": "paragraph", "text": line})

        i += 1

    return elements


def clean_markdown(text: str) -> str:
    """Remove markdown formatting from text."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)  # bold
    text = re.sub(r"\*(.+?)\*", r"\1", text)  # italic
    text = re.sub(r"`(.+?)`", r"\1", text)  # code
    text = text.replace("—", "-")  # em-dash to hyphen
    return text


def add_run_with_formatting(paragraph, text: str):
    """Add text to paragraph, handling **bold** inline."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            clean = clean_markdown(part)
            if clean:
                paragraph.add_run(clean)


def build_document(doc: Document, elements: list[dict], template: str = "plain"):
    """Populate document from parsed elements."""
    for el in elements:
        if el["type"] == "heading":
            p = doc.add_heading(clean_markdown(el["text"]), level=el["level"])

        elif el["type"] == "paragraph":
            p = doc.add_paragraph()
            add_run_with_formatting(p, el["text"])

        elif el["type"] == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            add_run_with_formatting(p, el["text"])

        elif el["type"] == "numbered":
            p = doc.add_paragraph(style="List Number")
            add_run_with_formatting(p, el["text"])

        elif el["type"] == "hr":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("_" * 60)
            run.font.color.rgb = RGBColor(180, 180, 180)
            run.font.size = Pt(8)

        elif el["type"] == "table":
            headers = el["headers"]
            rows = el["rows"]
            cols = max(len(headers), max((len(r) for r in rows), default=0))
            table = doc.add_table(rows=1 + len(rows), cols=cols)
            table.style = "Table Grid"

            # Header row
            for j, h in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = clean_markdown(h)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.size = Pt(11)

            # Data rows
            for i, row in enumerate(rows):
                for j, val in enumerate(row):
                    if j < cols:
                        cell = table.rows[i + 1].cells[j]
                        cell.text = clean_markdown(val)
                        for run in cell.paragraphs[0].runs:
                            run.font.size = Pt(11)

        elif el["type"] == "empty":
            pass  # skip empty lines (spacing handled by paragraph format)


def main():
    parser = argparse.ArgumentParser(description="DOCX generator (ГОСТ Р 7.0.97-2016)")
    parser.add_argument("--input", "-i", required=True, help="Input markdown file")
    parser.add_argument("--output", "-o", required=True, help="Output DOCX file")
    parser.add_argument("--template", "-t", default="plain",
                        choices=["plain", "kp", "letter"],
                        help="Document template (default: plain)")
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"ERROR: Input file not found: {input_path}")
        sys.exit(1)

    content = input_path.read_text(encoding="utf-8")
    elements = parse_markdown(content)

    doc = setup_document(args.template)
    build_document(doc, elements, args.template)

    doc.save(args.output)
    print(f"OK: {args.output} ({Path(args.output).stat().st_size} bytes)")


if __name__ == "__main__":
    main()
